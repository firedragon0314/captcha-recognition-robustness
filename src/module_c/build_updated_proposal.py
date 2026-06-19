from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parent
OUTPUT_PATH = PROJECT_ROOT / "Project_CO3027_吳睿嬛_葉俊廷_盛正璿_updated.docx"


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
    normal.font.size = Pt(11)

    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
        style.font.size = Pt(size)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(document: Document, headers, rows) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.add_paragraph()


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Programming for Deep Learning (CO3027)\nFinal Project Proposal")
    run.bold = True
    run.font.size = Pt(18)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Robust CAPTCHA Recognition under Corrupted Conditions:\n"
        "Sequence-based Recognition vs Position-wise Character Classification"
    )
    run.bold = True
    run.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Course: Programming for Deep Learning (CO3027)\n")
    p.add_run("Team: 吳睿嬛 / 葉俊廷 / 盛正璿")

    document.add_paragraph()


def build_document() -> Document:
    document = Document()
    set_document_defaults(document)
    add_title_page(document)

    document.add_heading("1. Topic", level=1)
    document.add_paragraph(
        "Robust CAPTCHA Recognition under Corrupted Conditions: "
        "Sequence-based Recognition vs Position-wise Character Classification"
    )

    document.add_heading("2. Brief Introduction of This Project", level=1)
    document.add_paragraph(
        "Text-based CAPTCHA is a common human verification mechanism, but its "
        "visual corruption patterns such as noise, blur, interference lines, rotation, "
        "occlusion, overlapping characters, and complex backgrounds make recognition difficult. "
        "This project studies CAPTCHA recognition in a controlled experimental setting and focuses "
        "on how different recognition strategies behave under clean, corrupted, and restored inputs."
    )
    document.add_paragraph(
        "The project compares two recognition paradigms and one restoration module:"
    )
    add_bullet(document, "Image restoration and data preparation")
    add_bullet(document, "Sequence-based CAPTCHA recognition")
    add_bullet(document, "Position-wise multi-head character classification")
    document.add_paragraph("Research questions:")
    add_bullet(document, "RQ1: Can image restoration improve recognition accuracy under corrupted conditions?")
    add_bullet(document, "RQ2: Which recognition strategy is more robust: CRNN-based sequence recognition or CNN multi-head character classification?")
    add_bullet(document, "RQ3: Can contrastive and self-supervised feature learning improve robustness under corruption?")

    document.add_heading("3. Related Work", level=1)

    document.add_heading("3.1 Image Restoration for Corrupted CAPTCHA", level=2)
    document.add_paragraph(
        "Image restoration is an important component of this project because corrupted CAPTCHA images often "
        "contain noise, blur, rotation, interference lines, and partial occlusion. For the first restoration "
        "branch, we refer to the Denoising Autoencoder (DAE) of Vincent et al. [1], which learns to "
        "reconstruct clean inputs from corrupted observations. This idea is directly relevant to our A-1 "
        "module, where corrupted CAPTCHA images are mapped back to cleaner versions through an encoder-decoder "
        "architecture."
    )
    document.add_paragraph(
        "For the second restoration branch, we adopt DnCNN [2], a residual-learning-based "
        "denoising model that predicts noise residuals instead of directly reconstructing the clean image. "
        "This design is closely related to our A-2 module, where restored CAPTCHA images are obtained by "
        "subtracting predicted residual noise from corrupted inputs."
    )
    document.add_paragraph(
        "Because our restoration experiments also evaluate structural image similarity, we further rely on "
        "SSIM [3] as an important metric for comparing clean, corrupted, and restored CAPTCHA images."
    )

    document.add_heading("3.2 Sequence-based CAPTCHA Recognition", level=2)
    document.add_paragraph(
        "For sequence-based recognition, our work is mainly based on the CRNN framework proposed by Shi, Bai, "
        "and Yao [5]. CRNN combines convolutional feature extraction with recurrent sequence modeling, making "
        "it suitable for unsegmented text recognition and overlapping characters. This is directly aligned with "
        "our B-1 module, where the whole CAPTCHA image is recognized as a sequence without explicit character segmentation."
    )
    document.add_paragraph(
        "To train sequence models without character-level alignment, we adopt Connectionist Temporal "
        "Classification (CTC) introduced by Graves et al. [4]. CTC is particularly suitable for our task "
        "because the model only needs the full CAPTCHA string label rather than the exact position of each character."
    )
    document.add_paragraph(
        "In addition to CRNN, our B-2 model incorporates a Transformer encoder for stronger global feature learning "
        "and long-range dependency modeling. For this part, we refer to the Vision Transformer work of "
        "Dosovitskiy et al. [6], which shows the effectiveness of transformer-based visual representations."
    )

    document.add_heading("3.3 Position-wise Multi-head Character Classification", level=2)
    document.add_paragraph(
        "For position-wise CAPTCHA recognition, our method follows a fixed-length multi-head classification design, "
        "where each output head predicts one character position independently. To strengthen the backbone in the "
        "improved version of our model, we refer to ResNet [7], which provides a strong and widely "
        "used convolutional feature extractor. This is especially relevant to our C-2 module, where a better CNN "
        "backbone is used for more robust feature learning."
    )
    document.add_paragraph(
        "Since CAPTCHA images may contain rotation, misalignment, and geometric distortion, we also refer to "
        "Spatial Transformer Networks [8]. STN is relevant because it provides a way to "
        "improve invariance to spatial transformation before character classification. This is useful for the "
        "improved position-wise model under corrupted conditions."
    )

    document.add_heading("3.4 Self-supervised and Semi-supervised Learning", level=2)
    document.add_paragraph(
        "To improve robustness under corruption, all three modules make use of self-supervised or contrastive "
        "learning ideas. Our main reference is SimCLR [9], which learns invariant feature "
        "representations by bringing augmented views of the same sample closer while pushing different samples "
        "apart. This directly supports our use of positive pairs, negative pairs, and augmentation consistency "
        "learning in A, B, and C."
    )

    document.add_heading("3.5 Optimization and Training Strategy", level=2)
    document.add_paragraph(
        "For optimization, we use Adam [10], because it is a standard and "
        "effective optimizer for deep learning training."
    )
    document.add_paragraph(
        "For learning-rate tuning, we refer to Cyclical Learning Rates [11], which highlights the "
        "importance of learning-rate search and scheduling in model performance. This is relevant to our "
        "experiments because we compare multiple learning-rate settings and analyze how learning-rate adjustment "
        "affects convergence and robustness."
    )

    document.add_heading("4. Dataset Specification", level=1)
    document.add_paragraph(
        "The CAPTCHA dataset is generated with the Python package "
        "`captcha.image.ImageCaptcha`. Each sample contains a fixed-length 5-character "
        "string composed of uppercase letters A-Z and digits 0-9."
    )
    add_table(
        document,
        ["Item", "Setting"],
        [
            ("Image size", "160 x 60"),
            ("CAPTCHA length", "5"),
            ("Character set", "A-Z and 0-9"),
            ("Random seed", "3027"),
            ("Dataset size", "100,000 clean CAPTCHA images"),
            ("Split", "70% train / 15% validation / 15% test"),
            ("Filename format", "index_text.png, e.g. 1023_A7K9P.png"),
        ],
    )
    document.add_paragraph("Data preparation stages:")
    add_number(document, "Generate clean CAPTCHA images.")
    add_number(document, "Create corrupted sets with Gaussian noise, motion blur, Gaussian blur, rotation, interference lines, background clutter, and partial occlusion.")
    add_number(document, "Apply restoration models to create restored CAPTCHA images for downstream recognition experiments.")

    document.add_heading("5. Team Project Plan / Methodology / Coding Architecture", level=1)
    document.add_paragraph(
        "The project is organized into three modules. The proposal below updates each module "
        "with the concrete technical scope described in A.docx, B.docx, and C.docx."
    )

    document.add_heading("5.1 Module A - Image Restoration and Data Preparation", level=2)
    document.add_paragraph(
        "This module is handled by 吳睿嬛. Its purpose is to generate corrupted CAPTCHA images, "
        "restore them with deep learning models, and provide clean/corrupted/restored datasets "
        "to Modules B and C for recognition experiments."
    )
    document.add_paragraph("Two restoration models will be implemented:")
    add_bullet(document, "Denoising Autoencoder (DAE): encoder-decoder reconstruction from corrupted CAPTCHA to clean CAPTCHA.")
    add_bullet(document, "DnCNN: residual-learning denoising model that predicts noise residuals and recovers restored CAPTCHA images.")
    document.add_paragraph("Corruption generation strategy:")
    add_bullet(document, "Gaussian noise with multiple sigma levels")
    add_bullet(document, "Blur with multiple kernel sizes")
    add_bullet(document, "Rotation with multiple severity levels")
    add_bullet(document, "Interference lines with multiple counts")
    add_bullet(document, "Partial occlusion with multiple ratios")
    document.add_paragraph("A will study multiple learning settings for restoration:")
    add_bullet(document, "Supervised learning on corrupted -> clean image pairs")
    add_bullet(document, "Unsupervised learning such as Noise2Noise or corruption-distribution learning")
    add_bullet(document, "Semi-supervised learning with limited clean pairs and consistency regularization")
    add_bullet(document, "Self-supervised learning such as contrastive learning or masked pixel prediction")
    document.add_paragraph("Key questions for Module A:")
    add_bullet(document, "Which restoration model produces the best restored CAPTCHA quality?")
    add_bullet(document, "Can restoration improve downstream recognition accuracy?")
    add_bullet(document, "Which corruption type is the hardest to repair?")
    add_table(
        document,
        ["Metric", "Description"],
        [
            ("Reconstruction Loss", "Training loss for restoration quality"),
            ("PSNR", "Peak Signal-to-Noise Ratio"),
            ("SSIM", "Structural Similarity Index"),
            ("Visual Comparison", "Side-by-side clean / corrupted / restored examples"),
        ],
    )

    document.add_heading("5.2 Module B - Sequence-based CAPTCHA Recognition", level=2)
    document.add_paragraph(
        "This module is handled by 葉俊廷. It recognizes the complete CAPTCHA string "
        "without manually segmenting characters, making it suitable for overlapping "
        "characters and irregular spacing."
    )
    document.add_paragraph("Module B includes two sequence models:")
    add_bullet(document, "B-1 CRNN with CTC Loss: CNN feature extractor -> sequence feature map -> RNN/LSTM/GRU -> CTC -> predicted CAPTCHA string.")
    add_bullet(document, "B-2 CNN + Transformer Encoder + CTC: sequence modeling with stronger long-range dependency learning.")
    document.add_paragraph("B will explore four learning settings:")
    add_bullet(document, "Supervised learning with image + full sequence label and CTC Loss")
    add_bullet(document, "Unsupervised contrastive learning to pretrain the CNN encoder with positive/negative CAPTCHA pairs")
    add_bullet(document, "Semi-supervised learning with labeled CAPTCHA and pseudo-labels on unlabeled data")
    add_bullet(document, "Self-supervised augmentation consistency learning under blur, noise, rotation, and brightness changes")
    document.add_paragraph("B focuses on:")
    add_bullet(document, "Sequence recognition accuracy")
    add_bullet(document, "Robustness under corrupted CAPTCHA conditions")
    add_bullet(document, "Recognition performance on restored CAPTCHA images")
    add_table(
        document,
        ["Model", "Train Data", "Test Data"],
        [
            ("Sequence CNN", "Clean", "Clean"),
            ("Sequence CNN", "Clean", "Corrupted"),
            ("Sequence CNN", "Restored", "Corrupted"),
            ("Sequence CNN", "Restored", "Restored"),
            ("Transformer OCR CNN", "Clean", "Clean"),
            ("Transformer OCR CNN", "Clean", "Corrupted"),
            ("Transformer OCR CNN", "Restored", "Corrupted"),
            ("Transformer OCR CNN", "Restored", "Restored"),
        ],
    )

    document.add_heading("5.3 Module C - Position-wise Multi-head CNN", level=2)
    document.add_paragraph(
        "This module is handled by 盛正璿. It treats CAPTCHA as a fixed-length string and "
        "predicts each character position independently with a shared encoder and multiple heads."
    )
    document.add_paragraph("Module C includes two position-wise variants:")
    add_bullet(document, "C-1 Pure Multi-head CNN: shared CNN encoder + 5 classification heads, each head performing 36-class prediction.")
    add_bullet(document, "C-2 Attention + Multi-head CNN: adds attention to focus on character regions and reduce background interference.")
    document.add_paragraph("Advantages of the multi-head design:")
    add_bullet(document, "Simple and stable training")
    add_bullet(document, "Fast inference")
    add_bullet(document, "Natural per-position analysis")
    add_bullet(document, "Suitable for fixed-length CAPTCHA")
    document.add_paragraph("Limitations to analyze:")
    add_bullet(document, "Sensitive to variable-length strings")
    add_bullet(document, "More affected by overlapping characters or severe position shifts")
    document.add_paragraph("C will explore four learning settings:")
    add_bullet(document, "Supervised learning with five position labels and summed cross-entropy loss")
    add_bullet(document, "Unsupervised contrastive learning for robust shared feature representation")
    add_bullet(document, "Semi-supervised learning with pseudo-labels for unlabeled CAPTCHA")
    add_bullet(document, "Self-supervised augmentation consistency learning for position-wise feature stability")
    document.add_paragraph("Additional engineering requirements for Module C:")
    add_bullet(document, "Track per-position accuracy, character accuracy, sequence accuracy, edit distance, robustness drop, and restoration gain")
    add_bullet(document, "Keep full experiment artifacts, including checkpoints, metrics, plots, and failure cases")
    add_bullet(document, "Record learning-rate changes and decision reasons during training")
    add_table(
        document,
        ["Metric", "Description"],
        [
            ("Per-position Accuracy", "Accuracy of each of the five character positions"),
            ("Character Accuracy", "Average accuracy across all characters"),
            ("Sequence Accuracy", "Whether the whole CAPTCHA is exactly correct"),
            ("Edit Distance", "Difference between predicted string and ground truth"),
            ("Robustness Drop", "Accuracy decrease after corruption"),
            ("Restoration Gain", "Accuracy increase after restoration"),
        ],
    )

    document.add_heading("6. Evaluation Metrics", level=1)
    document.add_paragraph(
        "The final evaluation combines restoration quality, recognition accuracy, and robustness analysis."
    )
    document.add_heading("6.1 Restoration Metrics", level=2)
    add_table(
        document,
        ["Metric", "Description"],
        [
            ("Reconstruction Loss", "Restoration training loss"),
            ("PSNR", "Pixel-level restoration quality"),
            ("SSIM", "Structural similarity"),
            ("Visual Comparison", "Clean / corrupted / restored examples"),
        ],
    )
    document.add_heading("6.2 Recognition Metrics", level=2)
    add_table(
        document,
        ["Metric", "Description"],
        [
            ("Character Accuracy", "Whether each predicted character is correct"),
            ("Sequence Accuracy", "Whether the full CAPTCHA string is exactly correct"),
            ("Edit Distance", "Distance between prediction and ground truth"),
            ("Per-position Accuracy", "Accuracy of positions 1 to 5 for Module C"),
        ],
    )
    document.add_heading("6.3 Robustness Metrics", level=2)
    add_table(
        document,
        ["Metric", "Description"],
        [
            ("Clean Accuracy", "Accuracy on the clean test set"),
            ("Noise Accuracy", "Accuracy on noisy CAPTCHA images"),
            ("Blur Accuracy", "Accuracy on blurred CAPTCHA images"),
            ("Rotation Accuracy", "Accuracy on rotated CAPTCHA images"),
            ("Interference-line Accuracy", "Accuracy on line-corrupted CAPTCHA images"),
            ("Restored Accuracy", "Accuracy after restoration"),
            ("Robustness Drop", "Clean Accuracy - Corrupted Accuracy"),
            ("Restoration Gain", "Restored Accuracy - Corrupted Accuracy"),
        ],
    )

    document.add_heading("7. Work Assignment and Five-Week Schedule", level=1)
    add_table(
        document,
        ["Member", "Main Responsibility", "AI / ML / DL Method"],
        [
            ("吳睿嬛", "CAPTCHA corruption generation and image restoration", "Denoising Autoencoder, DnCNN, restoration-oriented supervised / semi-supervised / self-supervised learning"),
            ("葉俊廷", "Sequence-based CAPTCHA recognition", "CRNN, Transformer Encoder, CTC Loss, Contrastive Learning, Semi-supervised and Self-supervised learning"),
            ("盛正璿", "Position-wise character classification and robustness learning", "Pure Multi-head CNN, Attention + Multi-head CNN, Contrastive Learning, Pseudo-labeling, Augmentation consistency learning"),
        ],
    )
    weeks = [
        ("Week 1", "Proposal discussion and project planning", "Finalize proposal, research questions, dataset plan, and division of responsibilities."),
        ("Week 2", "Dataset generation and baseline setup", "Generate clean/corrupted CAPTCHA data, build CRNN baseline, and build multi-head CNN baseline."),
        ("Week 3", "Main model implementation", "Implement DAE and DnCNN, train sequence models, train multi-head CNN models, and test feature pretraining."),
        ("Week 4", "Robustness evaluation and comparison", "Analyze restoration quality, evaluate sequence recognition, evaluate position-wise robustness, and compare methods."),
        ("Week 5", "Final integration and report", "Integrate results, prepare report/slides, finalize figures, code, and comparison tables."),
    ]
    add_table(document, ["Week", "Focus", "Expected Output"], weeks)

    document.add_heading("8. Final Deliverables", level=1)
    deliverables = [
        "A clean CAPTCHA generator together with corrupted and restored evaluation sets",
        "DAE and DnCNN restoration models with restoration-quality analysis",
        "CRNN and Transformer-based sequence-recognition models",
        "Pure Multi-head CNN and Attention + Multi-head CNN classifiers",
        "Evaluation tables under clean, corrupted, and restored conditions",
        "Robustness Drop and Restoration Gain analysis",
        "Stored checkpoints, metrics, plots, and failure-analysis records for reproducible experiments",
        "Final report and presentation slides",
    ]
    for item in deliverables:
        add_bullet(document, item)

    document.add_heading("9. References", level=1)
    references = [
        "[1] P. Vincent, H. Larochelle, Y. Bengio, and P.-A. Manzagol, \"Extracting and composing robust features with denoising autoencoders,\" in Proc. ICML, 2008.",
        "[2] K. Zhang, W. Zuo, Y. Chen, D. Meng, and L. Zhang, \"Beyond a Gaussian denoiser: Residual learning of deep CNN for image denoising,\" IEEE Trans. Image Process., vol. 26, no. 7, pp. 3142-3155, 2017.",
        "[3] Z. Wang, A. C. Bovik, H. R. Sheikh, and E. P. Simoncelli, \"Image quality assessment: From error visibility to structural similarity,\" IEEE Trans. Image Process., vol. 13, no. 4, pp. 600-612, 2004.",
        "[4] A. Graves, S. Fernandez, F. Gomez, and J. Schmidhuber, \"Connectionist temporal classification: Labelling unsegmented sequence data with recurrent neural networks,\" in Proc. ICML, 2006.",
        "[5] B. Shi, X. Bai, and C. Yao, \"An end-to-end trainable neural network for image-based sequence recognition and its application to scene text recognition,\" IEEE Trans. Pattern Anal. Mach. Intell., vol. 39, no. 11, pp. 2298-2304, 2017.",
        "[6] A. Dosovitskiy et al., \"An image is worth 16x16 words: Transformers for image recognition at scale,\" in Proc. ICLR, 2021.",
        "[7] K. He, X. Zhang, S. Ren, and J. Sun, \"Deep residual learning for image recognition,\" in Proc. CVPR, 2016.",
        "[8] M. Jaderberg, K. Simonyan, A. Zisserman, and K. Kavukcuoglu, \"Spatial transformer networks,\" in Proc. NeurIPS, 2015.",
        "[9] T. Chen, S. Kornblith, M. Norouzi, and G. Hinton, \"A simple framework for contrastive learning of visual representations,\" in Proc. ICML, 2020.",
        "[10] D. P. Kingma and J. Ba, \"Adam: A method for stochastic optimization,\" in Proc. ICLR, 2015.",
        "[11] L. N. Smith, \"Cyclical learning rates for training neural networks,\" in Proc. WACV, 2017.",
    ]
    for ref in references:
        document.add_paragraph(ref)

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"Saved {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
